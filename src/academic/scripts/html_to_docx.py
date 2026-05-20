#!/usr/bin/env python3
"""
HTML to DOCX converter with Chinese academic formatting support.

Converts HTML to DOCX with proper Chinese/English font separation,
following GB/T 7714-2015 academic formatting standards.

Usage:
    python html_to_docx.py input.html output.docx [options]

Options:
    --font-cn TEXT          Chinese font (default: SimSun)
    --font-en TEXT          English/number font (default: Times New Roman)
    --font-heading TEXT     Heading font (default: SimHei)
    --size-title INT        Title font size in pt (default: 16)
    --size-heading INT      Heading font size in pt (default: 15)
    --size-body INT         Body font size in pt (default: 14)
    --line-spacing FLOAT    Line spacing multiplier (default: 1.5)
"""

import argparse
import re
import sys
from pathlib import Path
from typing import Any

try:
    from bs4 import BeautifulSoup, Tag
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"Error: Missing required library: {e}", file=sys.stderr)
    print("Install with: pip install beautifulsoup4 python-docx lxml", file=sys.stderr)
    sys.exit(1)


def set_chinese_font(run, font_name_cn, font_name_en):
    """
    Set Chinese and English fonts separately.
    
    This solves the Word font separation issue where Chinese and English
    characters need different fonts.
    """
    run.font.name = font_name_en  # English font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)  # Chinese font


def add_paragraph_with_formatting(doc, text, style=None, font_cn='SimSun', 
                                   font_en='Times New Roman', font_size=14,
                                   bold=False, alignment=None):
    """Add a paragraph with proper Chinese/English font formatting."""
    para = doc.add_paragraph(style=style)
    
    if alignment:
        para.alignment = alignment
    
    run = para.add_run(text)
    set_chinese_font(run, font_cn, font_en)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    
    return para


def process_html_element(doc: Any, element: Tag, config: dict[str, Any]) -> None:
    """Process a single HTML element and add it to the document."""
    
    if element.name == 'h1':
        add_paragraph_with_formatting(
            doc, element.get_text().strip(),
            font_cn=config['font_heading'],
            font_en=config['font_en'],
            font_size=config['size_title'],
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
    
    elif element.name in ['h2', 'h3', 'h4', 'h5', 'h6']:
        add_paragraph_with_formatting(
            doc, element.get_text().strip(),
            font_cn=config['font_heading'],
            font_en=config['font_en'],
            font_size=config['size_heading'],
            bold=True
        )
    
    elif element.name == 'p':
        text = element.get_text().strip()
        if text:
            add_paragraph_with_formatting(
                doc, text,
                font_cn=config['font_cn'],
                font_en=config['font_en'],
                font_size=config['size_body']
            )
    
    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                para = add_paragraph_with_formatting(
                    doc, text,
                    style='List Bullet',
                    font_cn=config['font_cn'],
                    font_en=config['font_en'],
                    font_size=config['size_body']
                )
    
    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                para = add_paragraph_with_formatting(
                    doc, text,
                    style='List Number',
                    font_cn=config['font_cn'],
                    font_en=config['font_en'],
                    font_size=config['size_body']
                )
    
    elif element.name == 'blockquote':
        text = element.get_text().strip()
        if text:
            para = add_paragraph_with_formatting(
                doc, text,
                font_cn=config['font_cn'],
                font_en=config['font_en'],
                font_size=config['size_body'] - 1
            )
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
    
    elif element.name == 'table':
        process_table(doc, element, config)
    
    elif element.name == 'img':
        # Image handling (basic)
        src = element.get('src', '')
        alt = element.get('alt', '')
        if alt:
            add_paragraph_with_formatting(
                doc, f"[Image: {alt}]",
                font_cn=config['font_cn'],
                font_en=config['font_en'],
                font_size=config['size_body'] - 1,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )


def process_table(doc, table_element, config):
    """Process an HTML table and add it to the document."""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count columns
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                text = cell.get_text().strip()
                table_cell = table.rows[i].cells[j]
                para = table_cell.paragraphs[0]
                run = para.add_run(text)
                set_chinese_font(run, config['font_cn'], config['font_en'])
                run.font.size = Pt(config['size_body'] - 1)
                
                # Bold header cells
                if cell.name == 'th':
                    run.font.bold = True


def html_to_docx(html_path, docx_path, config):
    """
    Convert HTML to DOCX with Chinese academic formatting.
    
    Args:
        html_path: Path to input HTML file
        docx_path: Path to output DOCX file
        config: Dictionary with formatting options
    """
    
    # Read HTML
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create document
    doc = Document()
    
    # Set default line spacing
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Process body content
    body = soup.find('body')
    if body:
        for element in body.children:
            if isinstance(element, Tag):
                process_html_element(doc, element, config)
    
    # Apply line spacing to all paragraphs
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = config['line_spacing']
    
    # Save document
    doc.save(docx_path)


def main():
    parser = argparse.ArgumentParser(
        description='Convert HTML to DOCX with Chinese academic formatting'
    )
    parser.add_argument('input', help='Input HTML file')
    parser.add_argument('output', help='Output DOCX file')
    parser.add_argument('--font-cn', default='SimSun', 
                        help='Chinese font (default: SimSun)')
    parser.add_argument('--font-en', default='Times New Roman',
                        help='English/number font (default: Times New Roman)')
    parser.add_argument('--font-heading', default='SimHei',
                        help='Heading font (default: SimHei)')
    parser.add_argument('--size-title', type=int, default=16,
                        help='Title font size in pt (default: 16)')
    parser.add_argument('--size-heading', type=int, default=15,
                        help='Heading font size in pt (default: 15)')
    parser.add_argument('--size-body', type=int, default=14,
                        help='Body font size in pt (default: 14)')
    parser.add_argument('--line-spacing', type=float, default=1.5,
                        help='Line spacing multiplier (default: 1.5)')
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)
    
    # Build config
    config = {
        'font_cn': args.font_cn,
        'font_en': args.font_en,
        'font_heading': args.font_heading,
        'size_title': args.size_title,
        'size_heading': args.size_heading,
        'size_body': args.size_body,
        'line_spacing': args.line_spacing,
    }
    
    # Convert
    try:
        html_to_docx(args.input, args.output, config)
        print(f"Successfully converted {args.input} to {args.output}")
    except Exception as e:
        print(f"Error during conversion: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
